import os

import docx2txt
import pandas as pd
from docx import Document
from docx.shared import RGBColor


class Txt_to_xlsx:
    def __init__(self, user, separator, time):
        self.arr = []
        self.user = user
        self.time = time
        self.path = f"{self.user}.txt"
        self.path2 = f"{self.user}1.txt"
        self.path3 = f"{self.user}2.txt"
        self.excel = f"{self.user}.xlsx"
        self.docpath = f"{self.user}.docx"
        self.count = 0
        self.separator = separator
        self.store = f"Question Text{self.separator}Question Type{self.separator}Option 1{self.separator}Option 2{self.separator}Option 3{self.separator}Option 4{self.separator}Correct Answer{self.separator}Time in seconds{self.separator}Image Link\n"
        self.read_docx()
        self.reformattingfile()

        self.correctanswer()
        self.stringbuilding()

        # self.structurereplace()
        # self.csv_to_excel()
        os.remove(self.path)
        os.remove(self.path2)
        os.remove(self.path3)

    def read_docx(self):
        text = docx2txt.process(self.docpath)
        with open(self.path, "w", encoding='utf-8') as file:
            file.write(text)

    def correctanswer(self):
        for para in Document(self.docpath).paragraphs:
            print(para.runs)
            for run in para.runs:
                print(run)
            #     if (run.font.color.rgb == RGBColor(255, 000, 000)) or (run.underline) or (run.bold):
            #         if run.text.startswith("A."):
            #             temp = "1"
            #             self.arr.append(temp)

            #         elif run.text.startswith("B."):
            #             temp = "2"
            #             self.arr.append(temp)

            #         elif run.text.startswith("C."):
            #             temp = "3"
            #             self.arr.append(temp)

            #         elif run.text.startswith("D."):
            #             temp = "4"
            #             self.arr.append(temp)

    def reformattingfile(self):
        with open(self.path, "r") as file1:
            for line in file1:
                mystring = line.strip()
                for word in ["A. ", "B. ", "C. ", "D. "]:
                    mystring = mystring.replace(word, f"\n{word}")

                with open(self.path2, "a") as file2:
                    file2.write(mystring + "\n")

    def stringbuilding(self):
        with open(self.path2, "r") as file2:
            flag = False
            for line in file2:
                line = line.strip()

                if not line.startswith("A.") and flag:
                    self.store += line
                else:
                    if line.startswith("Câu "):
                        self.store += line
                        flag = True

                    elif line.startswith("A."):
                        flag = False
                        self.store += (self.separator + "Multiple Choice" + self.separator + line)

                    elif line.startswith("B.") or line.startswith("C."):
                        self.store += self.separator + line

                    elif line.startswith("D."):
                        # self.store += self.separator + line + self.separator + self.time + "\n"
                        self.store += self.separator + line + self.separator + self.arr[
                            self.count] + self.separator + self.time + "\n"
                        self.count += 1
                    # else:
                    #     self.store += line

    def structurereplace(self):
        with open(self.path3, 'a', encoding='utf-8') as file3:
            counter = 1
            while counter <= self.count:
                ques = f"Câu {counter}: "
                self.store = self.store.replace(ques, '')
                counter += 1

            counter = 1
            while counter <= self.count:
                ques = f"Câu {counter}. "
                self.store = self.store.replace(ques, '')
                counter += 1

            for char in ['A.', 'B.', 'C.', 'D.']:
                self.store = self.store.replace(char, '')

            file3.write(self.store)

    def debug(self):
        count = 0
        for para in Document(self.docpath).paragraphs:
            for run in para.runs:
                if (run.underline) or (run.font.color.rgb == RGBColor(255, 000, 000)) or (run.bold):
                    if run.text.startswith("A."):
                        temp = f"1-A-{count}"
                        self.arr.append(temp)
                
                    elif run.text.startswith("B."):
                        temp = f"2-B-{count}"
                        self.arr.append(temp)

                    elif run.text.startswith("C."):
                        temp = f"3-C-{count}"
                        self.arr.append(temp)

                    elif run.text.startswith("D."):
                        temp = f"4-D-{count}"
                        self.arr.append(temp)
        print(self.arr)
        print(len(self.arr))

    def csv_to_excel(self):
        csv = pd.read_csv(self.path3, sep=self.separator, on_bad_lines='skip')
        csv.to_excel(self.excel, index=False)


name = "/Users/jameshoang/Desktop/student fair/GV Bài 6"
c1 = Txt_to_xlsx(name, "|", "900")
